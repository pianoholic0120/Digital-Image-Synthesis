#!/usr/bin/env python3
"""Insert methods_comparison_6scenes.png into report.docx; renumber subsequent figures."""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

FINAL = Path(__file__).resolve().parents[2]
REPORT = FINAL / "report.docx"
FIGURE = FINAL / "report_figures" / "nerf_synthetic" / "methods_comparison_6scenes.png"

# Renumber old Figure 3–4 → 4–5 (order matters: longer patterns first).
FIGURE_RENUMBER = [
    (r"Figures 3-4", "Figures 4–5"),
    (r"Figures 3–4", "Figures 4–5"),
    (r"Figure 3–4", "Figure 4–5"),
    (r"Figure 3-4", "Figure 4–5"),
]


def renumber_figure_refs(text: str) -> str:
    for old, new in FIGURE_RENUMBER:
        text = text.replace(old, new)
    return text


def apply_renumbering(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip():
            p.text = renumber_figure_refs(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        p.text = renumber_figure_refs(p.text)


def find_paragraph(doc: Document, *, startswith: str) -> object | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    return None


def insert_block_before(anchor, *, image_path: Path, caption: str, body: str, width_in: float = 6.5) -> None:
    cap_p = anchor.insert_paragraph_before(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p = cap_p.insert_paragraph_before("")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))
    body_p = img_p.insert_paragraph_before(body)
    body_p.style = anchor.style


def main() -> int:
    if not REPORT.exists():
        print(f"Missing {REPORT}", file=sys.stderr)
        return 1
    if not FIGURE.exists():
        print(f"Missing {FIGURE}", file=sys.stderr)
        return 1

    doc = Document(str(REPORT))

    if find_paragraph(doc, startswith="Figure 3. Qualitative comparison on six NeRF Synthetic"):
        print("Figure 3 already present; skipping insert.")
        return 0

    # Shift Figure 3–4 → 4–5 before inserting new Figure 3.
    apply_renumbering(doc)

    anchor = find_paragraph(doc, startswith="Figures 4–5 show side-by-side")
    if anchor is None:
        anchor = find_paragraph(doc, startswith="Figures 4-5 show side-by-side")
    if anchor is None:
        anchor = find_paragraph(doc, startswith="Figure 4–5.")
    if anchor is None:
        anchor = find_paragraph(doc, startswith="Figure 4-5.")
    if anchor is None:
        sec = find_paragraph(doc, startswith="5.5  Qualitative Comparison")
        if sec is None:
            print("Could not find §5.5 anchor", file=sys.stderr)
            return 1
        idx = next(i for i, p in enumerate(doc.paragraphs) if p._element is sec._element)
        anchor = doc.paragraphs[idx + 1]

    body = (
        "Section 5.5 complements the quantitative metrics in Table 3 with side-by-side renders. "
        "Figure 3 summarizes qualitative agreement across six additional NeRF Synthetic test scenes "
        "(chair, drums, ficus, materials, mic, ship; 800×800, view 0). Each row compares stochastic "
        "ray tracing at 64 spp, our composite reference, the original 3DGS rasterizer, and the held-out "
        "ground-truth photograph. Stochastic and composite outputs are visually indistinguishable on these "
        "scenes, while differences against the rasterizer mirror the PSNR gaps in Table 3—for example, "
        "fine specular detail on chair and mic, and thin geometry on ship and ficus."
    )
    caption = (
        "Figure 3. Qualitative comparison on six NeRF Synthetic scenes (view 0): "
        "stochastic 64 spp, composite reference, 3DGS rasterizer, and ground truth."
    )

    insert_block_before(
        anchor,
        image_path=FIGURE,
        caption=caption,
        body=body,
        width_in=6.5,
    )

    # Revise hotdog/playroom lead-in to reference new Figure 3.
    intro = find_paragraph(doc, startswith="Figures 4–5 show side-by-side")
    if intro is None:
        intro = find_paragraph(doc, startswith="Figures 4-5 show side-by-side")
    if intro:
        intro.text = (
            "Figures 4–5 extend the qualitative comparison to SPP convergence on hotdog (NeRF Synthetic) "
            "and playroom (Tanks & Temples). Each panel shows stochastic ray tracing at increasing SPP "
            "(left columns) against the composite reference (right columns). At 32 spp and above the two "
            "images are visually indistinguishable, and both closely match the rasterizer output on coarse "
            "geometry. Fine structures are reproduced correctly by the ray tracer; we observe no systematic "
            "bias or splotching artifacts."
        )

    try:
        doc.save(str(REPORT))
        print(f"Updated {REPORT}")
    except PermissionError:
        alt = FINAL / "report_with_figure3.docx"
        doc.save(str(alt))
        print(f"Could not write {REPORT} (open in Word?). Saved {alt}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

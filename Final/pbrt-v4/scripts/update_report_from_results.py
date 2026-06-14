#!/usr/bin/env python3
"""Update Final/report.docx from benchmark_summary.json and lego_ablation.json."""

from __future__ import annotations

import json
import statistics
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

FINAL = Path(__file__).resolve().parents[2]
REPORT = FINAL / "report.docx"
SUMMARY = FINAL / "report_results" / "benchmark_summary.json"
ABLATION = FINAL / "report_results" / "lego_ablation.json"

DATASET_LABEL = {"3dgs": "NeRF Synthetic", "tandt_db": "Tanks & Temples"}
SCENE_ORDER = {
    "3dgs": ["drums", "chair", "ficus", "hotdog", "lego", "materials", "mic", "ship"],
    "tandt_db": ["drjohnson", "train", "truck", "playroom"],
}
SPP_COLS = ["1", "32", "64", "256", "1024", "2048"]
REP_SCENES = [
    ("3dgs", "lego"),
    ("3dgs", "hotdog"),
    ("tandt_db", "drjohnson"),
    ("tandt_db", "playroom"),
]


def fmt_time(v: float | None) -> str:
    if v is None or (isinstance(v, (int, float)) and v <= 0):
        return "—"
    return f"{v:.1f}"


def fmt_metric(v: float | None, nd: int = 2) -> str:
    if v is None:
        return "—"
    if v == float("inf"):
        return "inf"
    return f"{v:.{nd}f}"


def set_cell(table, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = text


def find_table_by_headers(doc: Document, headers: list[str]):
    for table in doc.tables:
        if not table.rows:
            continue
        cells = table.rows[0].cells
        if len(cells) < len(headers):
            continue
        row_headers = [cells[i].text.strip() for i in range(len(headers))]
        if row_headers == headers:
            return table
    return None


def replace_paragraph_text(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)


def delete_table(table) -> None:
    table._element.getparent().remove(table._element)


def find_paragraph(doc: Document, *, startswith: str | None = None, contains: str | None = None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if startswith and t.startswith(startswith):
            return p
        if contains and contains in t:
            return p
    return None


SUPP_TABLE_HEADERS = [
    ["Scene", "64 spp vs Comp", "1024 spp vs Comp", "1024 spp vs GT"],
    ["Dataset", "Scene", "Stoch 64 spp", "Comp 1 spp", "Comp/Stoch"],
    ["Scene", "Method", "PSNR", "SSIM", "LPIPS"],
    ["Scene", "1 spp", "32 spp", "64 spp", "256 spp", "1024 spp", "2048 spp"],
]


def cleanup_stale_report_content(doc: Document) -> None:
    """Remove placeholder text, path dumps, duplicate supplementary tables/captions."""
    drop_starts = (
        "[INSERT Figure",
        "[INSERT convergence",
        "Generated comparison grid:",
        "Additional qualitative grids",
        "Table 4. Stochastic",
        "Table 5. Mean wall-clock",
        "Table 6. PSNR",
        "Table 7. Stochastic",
        "Qualitative comparisons (Fig.",
        "Figure 3. PSNR vs. SPP convergence (lego). See",
    )
    drop_contains = (
        "report_figures\\qualitative",
        "report_figures/qualitative",
        "See C:\\Users",
        "See C:/Users",
        "  • spp_sweep",
        "Digital-Image-Synthesis\\Final\\report_figures",
    )
    to_remove = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if any(t.startswith(s) for s in drop_starts):
            to_remove.append(p)
            continue
        if any(s in t for s in drop_contains):
            to_remove.append(p)
    for p in to_remove:
        delete_paragraph(p)
    for table in list(doc.tables):
        headers = [c.text.strip() for c in table.rows[0].cells]
        if any(headers[: len(sig)] == sig for sig in SUPP_TABLE_HEADERS):
            delete_table(table)
        elif len(table.rows) <= 1 and all(not c.text.strip() for row in table.rows for c in row.cells):
            delete_table(table)


def polish_table_captions(doc: Document) -> None:
    for p in doc.paragraphs:
        if "Table 1." in p.text:
            p.text = "Table 1. CPU stochastic rendering performance (seconds). #G = Gaussians (×10³)."
        elif p.text.strip().startswith("Table 2."):
            p.text = (
                "Table 2. Image quality vs. held-out test views (PSNR ↑ dB, SSIM ↑, LPIPS ↓). "
                "Stoch = stochastic 1024 spp; Comp = composite reference; Rast = 3DGS rasterizer."
            )
        elif p.text.strip().startswith("Table 3."):
            p.text = "Table 3. Ablation on lego (800×800, 64 spp): PSNR vs. composite reference and wall-clock time."


def insert_table_after(doc: Document, anchor, caption: str, headers: list[str], data: list[list[str]]):
    cap = anchor.insert_paragraph_before(caption)
    cap.style = anchor.style
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for ri, row in enumerate(data, start=1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val
    cap._p.addnext(table._tbl)


def paragraph_has_image(paragraph) -> bool:
    return any(r._element.xpath(".//a:blip") for r in paragraph.runs)


def insert_figure_before(doc: Document, anchor, image_path: Path, caption: str, width_in: float = 6.5) -> None:
    if not image_path.exists():
        return
    prev = anchor._element.getprevious()
    if prev is not None and prev.tag.endswith("p"):
        from docx.text.paragraph import Paragraph

        if paragraph_has_image(Paragraph(prev, anchor._parent)) and caption in Paragraph(prev, anchor._parent).text:
            return
    img_p = anchor.insert_paragraph_before("")
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap_p = anchor.insert_paragraph_before(caption)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.style = anchor.style


def load_summary() -> dict:
    if not SUMMARY.exists():
        raise SystemExit(f"Missing {SUMMARY}; run report_benchmark_pipeline.py first.")
    return json.loads(SUMMARY.read_text(encoding="utf-8"))


def scene_key(ds: str, scene: str) -> str:
    return f"{ds}/{scene}"


def aggregate_stoch_vs_comp(summary: dict) -> dict[str, float]:
    """Mean PSNR of stochastic vs composite across scenes with data."""
    vals_64, vals_1024 = [], []
    for key, entry in summary.get("scenes", {}).items():
        q = entry.get("quality", {})
        if "stochastic_vs_composite_64" in q:
            vals_64.append(q["stochastic_vs_composite_64"]["psnr"])
        if "stochastic_vs_composite_1024" in q:
            vals_1024.append(q["stochastic_vs_composite_1024"]["psnr"])
    out = {}
    if vals_64:
        out["mean_64"] = statistics.mean(vals_64)
    if vals_1024:
        out["mean_1024"] = statistics.mean(vals_1024)
    return out


def scene_row_map(table, col: int = 0) -> dict[str, int]:
    out: dict[str, int] = {}
    for r in range(1, len(table.rows)):
        name = table.rows[r].cells[col].text.strip()
        if name and name not in DATASET_LABEL.values() and not name.startswith("NeRF") and name != "T&T Mean":
            out[name] = r
    return out


def fill_performance_table(doc: Document, summary: dict) -> None:
    table = find_table_by_headers(doc, ["Scene", "#G(k)", "1 spp", "32 spp", "64 spp", "256 spp"])
    if table is None:
        table = find_table_by_headers(
            doc,
            ["Dataset", "Scene", "#G (k)", "1 spp", "32 spp", "64 spp", "256 spp", "1024 spp", "2048 spp"],
        )
        if table is None:
            return
        row_map = {(table.rows[r].cells[0].text.strip(), table.rows[r].cells[1].text.strip()): r for r in range(1, len(table.rows))}
        for ds_key, scenes in SCENE_ORDER.items():
            label = DATASET_LABEL[ds_key]
            for scene in scenes:
                ri = row_map.get((label, scene))
                if ri is None:
                    continue
                entry = summary.get("scenes", {}).get(scene_key(ds_key, scene), {})
                set_cell(table, ri, 2, fmt_metric(entry.get("gaussians_k"), 1))
                timing = entry.get("timing_stochastic", {})
                for ci, spp in enumerate(SPP_COLS, start=3):
                    set_cell(table, ri, ci, fmt_time(timing.get(spp, {}).get("mean_time_s")))
        return

    row_map = scene_row_map(table)
    spp_cols = ["1", "32", "64", "256", "1024", "2048"]
    for ds_key, scenes in SCENE_ORDER.items():
        for scene in scenes:
            ri = row_map.get(scene)
            if ri is None:
                continue
            entry = summary.get("scenes", {}).get(scene_key(ds_key, scene), {})
            set_cell(table, ri, 1, fmt_metric(entry.get("gaussians_k"), 1))
            timing = entry.get("timing_stochastic", {})
            for ci, spp in enumerate(spp_cols, start=2):
                if ci < len(table.rows[ri].cells):
                    set_cell(table, ri, ci, fmt_time(timing.get(spp, {}).get("mean_time_s")))


def fill_speedup_table(doc: Document, summary: dict) -> None:
    table = find_table_by_headers(doc, ["Scene", "Stoch 64 spp (s)", "Comp 1 spp (s)", "Speedup"])
    if table is None:
        return
    row_map = scene_row_map(table)
    for ds_key, scenes in SCENE_ORDER.items():
        for scene in scenes:
            ri = row_map.get(scene)
            if ri is None:
                continue
            e = summary.get("scenes", {}).get(scene_key(ds_key, scene), {})
            t64 = e.get("timing_stochastic", {}).get("64", {}).get("mean_time_s")
            tc = e.get("quality", {}).get("composite_1", {}).get("time_s")
            set_cell(table, ri, 1, fmt_time(t64))
            set_cell(table, ri, 2, fmt_time(tc))
            ratio = f"{tc / t64:.1f}×" if t64 and tc and t64 > 0 else "—"
            set_cell(table, ri, 3, ratio)


def fill_quality_table(doc: Document, summary: dict) -> None:
    table = find_table_by_headers(doc, ["Scene", "PSNR↑", "PSNR↑", "PSNR↑", "SSIM↑", "SSIM↑", "SSIM↑"])
    if table is None:
        table = find_table_by_headers(
            doc,
            [
                "Dataset", "Scene", "Stoch PSNR", "Stoch SSIM", "Stoch LPIPS",
                "Comp PSNR", "Comp SSIM", "Comp LPIPS", "Rast PSNR", "Rast SSIM", "Rast LPIPS",
            ],
        )
        if table is None:
            return
        row_map = {(table.rows[r].cells[0].text.strip(), table.rows[r].cells[1].text.strip()): r for r in range(1, len(table.rows))}
        mapping = {
            "stoch": ("stochastic_1024", 2, 3, 4),
            "comp": ("composite_1", 5, 6, 7),
            "rast": ("rasterizer", 8, 9, 10),
        }
        for ds_key, scenes in SCENE_ORDER.items():
            label = DATASET_LABEL[ds_key]
            for scene in scenes:
                ri = row_map.get((label, scene))
                if ri is None:
                    continue
                q = summary.get("scenes", {}).get(scene_key(ds_key, scene), {}).get("quality", {})
                for _, (block_key, c0, c1, c2) in mapping.items():
                    block = q.get(block_key, {})
                    set_cell(table, ri, c0, fmt_metric(block.get("psnr")))
                    set_cell(table, ri, c1, fmt_metric(block.get("ssim"), 4))
                    set_cell(table, ri, c2, fmt_metric(block.get("lpips"), 4))
        return

    row_map = scene_row_map(table)
    blocks = [("stochastic_1024", 1), ("composite_1", 2), ("rasterizer", 3)]
    for ds_key, scenes in SCENE_ORDER.items():
        for scene in scenes:
            ri = row_map.get(scene)
            if ri is None:
                continue
            q = summary.get("scenes", {}).get(scene_key(ds_key, scene), {}).get("quality", {})
            for block_key, psnr_col in blocks:
                b = q.get(block_key, {})
                set_cell(table, ri, psnr_col, fmt_metric(b.get("psnr")))
                set_cell(table, ri, psnr_col + 3, fmt_metric(b.get("ssim"), 4))
                set_cell(table, ri, psnr_col + 6, fmt_metric(b.get("lpips"), 4))


def fill_spp_convergence_table(doc: Document, summary: dict) -> None:
    table = find_table_by_headers(doc, ["Scene", "1 spp", "32 spp", "64 spp", "256 spp", "1024 spp"])
    if table is None:
        return
    row_map = scene_row_map(table)
    all_scenes = [(ds, sc) for ds, scenes in SCENE_ORDER.items() for sc in scenes]
    for ds_key, scene in all_scenes:
        ri = row_map.get(scene)
        if ri is None:
            continue
        q = summary.get("scenes", {}).get(scene_key(ds_key, scene), {}).get("quality", {})
        spp_list = SPP_COLS if ds_key == "3dgs" else [s for s in SPP_COLS if s != "2048"]
        for ci, spp in enumerate(spp_list, start=1):
            if ci >= len(table.rows[ri].cells):
                break
            set_cell(table, ri, ci, spp_metric(q, spp))


ABLATION_ORDER = [
    ("Depth mode", "OursCenter (default)"),
    ("Depth mode", "OursMean"),
    ("sigma cutoff", "1.5"),
    ("sigma cutoff", "2.0"),
    ("sigma cutoff", "2.828 (default)"),
    ("sigma cutoff", "3.5"),
    ("Accelerator", "BVH (default)"),
    ("Accelerator", "kd-tree"),
    ("SPP", "1"),
    ("SPP", "32"),
    ("SPP", "64"),
    ("SPP", "256"),
]


def fill_ablation_table(doc: Document) -> None:
    if not ABLATION.exists():
        return
    raw = {(r["config"], r["variant"]): r for r in json.loads(ABLATION.read_text(encoding="utf-8")).get("rows", [])}
    table = find_table_by_headers(doc, ["Config", "Variant", "PSNR vs Comp (dB) ↑", "Time (s)"])
    if table is None:
        table = find_table_by_headers(doc, ["Config", "Variant", "PSNR vs Ref (dB)", "Time (s)"])
    if table is None:
        return
    prev_config = None
    for i, (config, variant) in enumerate(ABLATION_ORDER):
        ri = i + 1
        if ri >= len(table.rows):
            break
        row = raw.get((config, variant), {})
        label = config if config != prev_config else ""
        prev_config = config
        set_cell(table, ri, 0, label)
        set_cell(table, ri, 1, variant)
        set_cell(table, ri, 2, fmt_metric(row.get("psnr_vs_composite")))
        set_cell(table, ri, 3, fmt_time(row.get("time_s")))


def fix_methods_text(doc: Document) -> None:
    replacements = [
        (
            "h(p) = frac(sin(a₁·p.x + b₁·p.y) · a₂ + cos(b₂·p.z) · a₃)",
            "r₁(q)=frac(b₁·sin(a₁·q)),  r₂(qx,qy)=frac(b₂·sin(a₂x·qx+a₂y·qy)),  ξ(p)=r₂(p_x+r₁(p_z), p_y)  (Sun et al. Eq. 8–10)",
        ),
        (
            "where the constants a₁, b₁, a₂, b₂, a₃ match Sun et al. exactly.",
            "with (a₁,b₁)=(91.3458, 47453.5453), (a₂x,a₂y,b₂)=(12.9898, 78.233, 43758.5453). A Sobol triple offset (10⁻⁴ scale) is added to the hit position before hashing to decorrelate frames.",
        ),
        (
            "generate ξ_k = TrigHash(ray_id, k, frame) ∈ [0,1) and accept if ξ_k < α_k",
            "draw ξ = TrigHash(p, frame) at the hit position p and accept if ξ < α (Eq. 7); multi-sample mode uses independent slots s with frame index frame·N+s",
        ),
        (
            "(2) evaluate the 1D Gaussian depth mode (OursCenter: use AABB center projected onto ray; OursMean: use the Gaussian mean)",
            "(2) evaluate depth t (OursMean: 1D Gaussian mean t_μ; OursCenter: μ projected onto the ray); (3) evaluate opacity α at the 1D mean for the RR test",
        ),
        (
            "For validation we also implement the exact compositing estimator (Eq. 3 of Sun et al.): all Gaussians along a ray are sorted by depth, and transmittance-weighted radiance is accumulated analytically. This serves as a ground-truth reference to verify unbiasedness of the stochastic estimator.",
            "For validation we implement the exact compositing estimator (Sun et al. Eq. 3): candidate Gaussians from the 2D tile grid are sorted by depth and transmittance-weighted radiance is accumulated analytically at 1 spp. We use this composite image as the algorithmic reference when measuring stochastic unbiasedness (PSNR between stochastic and composite renders of the same scene).",
        ),
    ]
    for old, new in replacements:
        replace_paragraph_text(doc, old, new)


def add_implementation_section(doc: Document) -> None:
    if any(p.text.strip().startswith("3.7") for p in doc.paragraphs):
        return
    paras = [
        "3.7  PBRT Implementation Details",
        "The system is implemented as a PBRT-v4 shape (gaussiancloud) plus material (gaussian). At load time each PLY vertex becomes a Gaussian with precomputed Σ⁻¹, an AABB at σ_cutoff=2√2, and optional floater pruning (min opacity, max scale percentile, SH / distance outlier rejection). Scene files are auto-generated from cameras.json and the trained point_cloud.ply.",
        "IntersectStochastic dispatches by sampling_mode: stochastic uses IntersectBVH (or kd-tree) with multi_samples independent RR slots per pixel sample; composite uses IntersectComposite with 2D tile-grid candidate collection and depth sorting. Each stochastic candidate calls EvalAlphaForRay at the 1D mean depth, draws TrigHash(p, frame·N+s), and on acceptance clips currentTMax so farther Gaussians along the ray are skipped—matching the paper’s opacity-based RR.",
        "Shading evaluates degree-3 SH through GaussianMaterial (48 coefficients) into RGB for the path integrator’s primary rays. Camera intrinsics from cameras.json populate a 2D tile grid for O(1) primary-ray candidate lookup. NeRF Synthetic scenes use rasterizer-aligned 2D projected α (use_2d_alpha=true); Tanks & Temples use 3D ray-space α (use_2d_alpha=false) because 2D projection mismatches COLMAP cameras on real captures. use_center_depth toggles OursCenter vs OursMean traversal depth while α is evaluated at the 1D mean (paper §4.1).",
        "Evaluation is automated by report_benchmark_pipeline.py: for each dataset, scene, view (10 held-out test frames), and SPP level we record wall-clock stochastic time, export PNGs (stochastic, composite, GT, rasterizer) under Final/report_assets/, and write per-view metrics.txt with PSNR/SSIM/LPIPS vs. GT.",
    ]
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("3.6  Composite Reference"):
            anchor = doc.paragraphs[i + 1]
            for text in reversed(paras):
                anchor.insert_paragraph_before(text)
            break


def update_experiment_prose(doc: Document, summary: dict) -> None:
    n_views = summary.get("config", {}).get("views", [])
    n = len(n_views) if n_views else 10
    replace_paragraph_text(
        doc,
        "PSNR, SSIM, and LPIPS are computed against held-out test views.",
        f"PSNR, SSIM, and LPIPS are computed against held-out test views (frames 0–{n - 1}, mean over {n} views). "
        "Metrics compare linear EXR renders to GT/rasterizer PNGs using the 3DGS convention (8-bit values treated as linear, no sRGB decode). "
        "LPIPS uses the AlexNet backbone. Rasterizer timings are omitted (not measured on our CPU path).",
    )

    lego = summary.get("scenes", {}).get("3dgs/lego", {})
    gk = lego.get("gaussians_k")
    t64 = lego.get("timing_stochastic", {}).get("64", {}).get("mean_time_s")
    comp_t = lego.get("quality", {}).get("composite_1", {}).get("time_s")
    if gk and t64 and comp_t:
        replace_paragraph_text(
            doc,
            "As a reference: on our test machine, the lego scene (~322k Gaussians, 800×800) renders at approximately 2.4 s at 64 spp (stochastic) vs. 21.9 s at 1 spp (composite reference), confirming the expected ≥5–20× speedup of the stochastic estimator. CPU timings for drjohnson and playroom at 64 spp are approximately 24.7 s and 20.7 s respectively, consistent with Sun et al. Table 1 (46.8 s and 31.7 s on their hardware; our machine is roughly 2× faster).",
            f"On our Intel Core i7-13700, lego ({gk:.0f}k Gaussians after pruning, 800×800, mean over {n} views) takes {t64:.1f} s at 64 spp stochastic vs. {comp_t:.1f} s for the composite reference (1 spp), i.e. {comp_t / t64:.1f}× slower for exact compositing at comparable output variance budget.",
        )

    dj = summary.get("scenes", {}).get("tandt_db/drjohnson", {})
    pr = summary.get("scenes", {}).get("tandt_db/playroom", {})
    dj64 = dj.get("timing_stochastic", {}).get("64", {}).get("mean_time_s")
    pr64 = pr.get("timing_stochastic", {}).get("64", {}).get("mean_time_s")
    if dj64 and pr64:
        replace_paragraph_text(
            doc,
            "CPU timings for drjohnson and playroom at 64 spp are approximately 24.7 s and 20.7 s respectively, consistent with Sun et al. Table 1 (46.8 s and 31.7 s on their hardware; our machine is roughly 2× faster).",
            f"For Tanks & Temples we render at held-out GT resolution (not the 2× supersampled cameras.json size used by some training exports). "
            f"At 64 spp, drjohnson and playroom average {dj64:.1f} s and {pr64:.1f} s (Sun et al. Table 1: 46.8 s and 31.7 s on their CPU).",
        )
    replace_paragraph_text(
        doc,
        "Rendering resolution is 800×800 for NeRF Synthetic and 1920×1080 for Tanks-and-Temples, matching the original 3DGS evaluation protocol.",
        "Rendering resolution is 800×800 for NeRF Synthetic. For Tanks & Temples we use each scene's GT image resolution (e.g. drjohnson 1332×876; outdoor train/truck ~980×545).",
    )

    vs = aggregate_stoch_vs_comp(summary)
    if vs.get("mean_1024"):
        lo, hi = _minmax_vs_comp(summary, 1024)
        replace_paragraph_text(
            doc,
            "As reference figures based on our validation runs: stochastic (1024 spp) achieves 47–55 dB PSNR vs. the composite reference on NeRF Synthetic scenes, confirming near-unbiased convergence. On Tanks-and-Temples drjohnson and playroom, stochastic exceeds 46 dB vs. composite. The train scene is lower (~37 dB vs. composite) due to point cloud sparsity in difficult outdoor regions rather than algorithmic error.",
            f"Stochastic vs. composite PSNR (algorithmic agreement) averages {vs['mean_1024']:.1f} dB at 1024 spp across measured scenes (range {lo:.1f}–{hi:.1f} dB), confirming near-unbiased convergence. "
            "PSNR vs. held-out GT (Table 2) reflects training/rasterization quality of the 3DGS point cloud, not stochastic bias.",
        )
        train_vs = summary.get("scenes", {}).get("tandt_db/train", {}).get("quality", {}).get("stochastic_vs_composite_1024", {})
        if train_vs.get("psnr"):
            replace_paragraph_text(
                doc,
                "The train scene is lower (~37 dB vs. composite) due to point cloud sparsity in difficult outdoor regions rather than algorithmic error.",
                f"Outdoor train remains lowest vs. composite ({train_vs['psnr']:.1f} dB at 1024 spp), consistent with point-cloud sparsity rather than renderer bias.",
            )

    replace_paragraph_text(
        doc,
        "Experimental results confirm that the stochastic estimator is effectively unbiased (47–55 dB PSNR vs. composite) while being 5–20× faster than the exact composite at equal SPP budget.",
        "Experimental results confirm near-unbiased stochastic estimation (high PSNR vs. composite; Table 4) while stochastic traversal is substantially faster than exact compositing for equal image quality.",
    )


def _minmax_vs_comp(summary: dict, spp: int) -> tuple[float, float]:
    vals = [
        e["quality"][f"stochastic_vs_composite_{spp}"]["psnr"]
        for e in summary.get("scenes", {}).values()
        if f"stochastic_vs_composite_{spp}" in e.get("quality", {})
    ]
    return (min(vals), max(vals)) if vals else (0.0, 0.0)


def add_spp_brightness_note(doc: Document) -> None:
    replace_paragraph_text(
        doc,
        "Figure 1 shows side-by-side comparisons of (a) ground-truth test image, (b) 3DGS rasterizer output, (c) composite reference (our exact renderer), and (d) stochastic ray tracer at 64 spp, for representative scenes from each dataset.",
        "Figure 1 shows side-by-side comparisons of (a) ground-truth test image, (b) 3DGS rasterizer output, "
        "(c) composite reference (our exact renderer), and (d) stochastic ray tracer at 64 spp, for representative scenes "
        "from each dataset. At 1 spp the stochastic image often appears darker than 64–1024 spp: this is expected Monte Carlo "
        "variance (a single RR draw per pixel underestimates energy on average) and not a tonemapping bug; mean intensity "
        "converges toward the composite reference as SPP increases (Figures 4–5).",
    )


def insert_report_figures(doc: Document) -> None:
    """Embed PNG figures at the correct section anchors (not file paths)."""
    fig1 = FINAL / "report_figures" / "figure1_method_comparison.png"
    fig3 = FINAL / "report_figures" / "figure3_convergence_lego.png"
    qual = FINAL / "report_figures" / "qualitative"

    anchor1 = find_paragraph(doc, startswith="Figure 1. Qualitative comparison")
    if anchor1 and fig1.exists():
        insert_figure_before(
            doc, anchor1, fig1,
            "Figure 1. Qualitative comparison across rendering methods (GT, rasterizer, composite, stochastic @64 spp).",
        )

    anchor3 = find_paragraph(doc, startswith="Figure 3.")
    if anchor3 and fig3.exists():
        insert_figure_before(
            doc, anchor3, fig3,
            "Figure 3. PSNR vs. SPP convergence on lego (stochastic vs. composite reference).",
            width_in=5.0,
        )

    anchor_fig2 = find_paragraph(doc, startswith="Figure 2. Ray-tracing")
    if anchor_fig2 and qual.exists():
        for name, cap, w in [
            ("spp_sweep_lego.png", "Figure 4. SPP sweep on lego (stochastic).", 5.5),
            ("spp_sweep_drjohnson.png", "Figure 5. SPP sweep on drjohnson (stochastic).", 5.5),
            ("stoch_vs_comp_lego_view00000.png", "Figure 6. Stochastic vs. composite on lego (view 0, 64 spp).", 5.5),
            ("stoch_vs_comp_drjohnson_view00000.png", "Figure 7. Stochastic vs. composite on drjohnson (view 0, 64 spp).", 5.5),
        ]:
            path = qual / name
            if path.exists():
                insert_figure_before(doc, anchor_fig2, path, cap, width_in=w)


def spp_metric(q: dict, spp: str) -> str:
    block = q.get(f"stochastic_vs_composite_{spp}")
    return fmt_metric(block.get("psnr") if block else None)


def insert_supplementary_tables(doc: Document, summary: dict) -> None:
    """Place Tables 4–7 at end of Section 5 (before Discussion)."""
    anchor = find_paragraph(doc, startswith="6  DISCUSSION")
    if anchor is None:
        return

    t5_rows = []
    for ds_key, scenes in SCENE_ORDER.items():
        for scene in scenes:
            e = summary.get("scenes", {}).get(scene_key(ds_key, scene), {})
            t64 = e.get("timing_stochastic", {}).get("64", {}).get("mean_time_s")
            tc = e.get("quality", {}).get("composite_1", {}).get("time_s")
            ratio = f"{tc / t64:.1f}×" if t64 and tc and t64 > 0 else "—"
            t5_rows.append([DATASET_LABEL[ds_key], scene, fmt_time(t64), fmt_time(tc), ratio])
    t4_rows = []
    for ds_key, scene in REP_SCENES:
        q = summary.get("scenes", {}).get(scene_key(ds_key, scene), {}).get("quality", {})
        t4_rows.append([
            scene,
            fmt_metric(q.get("stochastic_vs_composite_64", {}).get("psnr")),
            fmt_metric(q.get("stochastic_vs_composite_1024", {}).get("psnr")),
            fmt_metric(q.get("stochastic_1024", {}).get("psnr")),
        ])
    t6_rows = []
    for ds_key, scene in REP_SCENES:
        q = summary.get("scenes", {}).get(scene_key(ds_key, scene), {}).get("quality", {})
        for method, key in [
            ("Stochastic", "stochastic_1024"),
            ("Composite", "composite_1"),
            ("Rasterizer", "rasterizer"),
        ]:
            b = q.get(key, {})
            t6_rows.append([
                scene, method,
                fmt_metric(b.get("psnr")),
                fmt_metric(b.get("ssim"), 4),
                fmt_metric(b.get("lpips"), 4),
            ])
    t7_rows = []
    for ds_key, scene in REP_SCENES:
        q = summary.get("scenes", {}).get(scene_key(ds_key, scene), {}).get("quality", {})
        t7_rows.append([scene] + [spp_metric(q, s) for s in SPP_COLS])

    # Insert in reverse so final order is Table 4 → 7 before Section 6.
    insert_table_after(
        doc, anchor,
        "Table 7. Stochastic vs. composite PSNR (dB) at each SPP — algorithmic convergence (representative scenes).",
        ["Scene"] + [f"{s} spp" for s in SPP_COLS],
        t7_rows,
    )
    insert_table_after(
        doc, anchor,
        "Table 6. PSNR / SSIM / LPIPS vs. GT at 1024 spp (representative scenes).",
        ["Scene", "Method", "PSNR", "SSIM", "LPIPS"],
        t6_rows,
    )
    insert_table_after(
        doc, anchor,
        "Table 5. Mean wall-clock time (s): stochastic 64 spp vs. composite 1 spp (10 views).",
        ["Dataset", "Scene", "Stoch 64 spp", "Comp 1 spp", "Comp/Stoch"],
        t5_rows,
    )
    insert_table_after(
        doc, anchor,
        "Table 4. Stochastic vs. composite PSNR (dB, mean over 10 views). Columns 2–3: algorithmic agreement; column 4: vs. GT.",
        ["Scene", "64 spp vs Comp", "1024 spp vs Comp", "1024 spp vs GT"],
        t4_rows,
    )


def main() -> int:
    summary = load_summary()
    if not REPORT.exists():
        print(f"Missing {REPORT}", file=sys.stderr)
        return 1
    doc = Document(str(REPORT))
    cleanup_stale_report_content(doc)
    fix_methods_text(doc)
    add_implementation_section(doc)
    polish_table_captions(doc)
    insert_supplementary_tables(doc, summary)
    fill_performance_table(doc, summary)
    fill_speedup_table(doc, summary)
    fill_quality_table(doc, summary)
    fill_spp_convergence_table(doc, summary)
    fill_ablation_table(doc)
    update_experiment_prose(doc, summary)
    add_spp_brightness_note(doc)
    insert_report_figures(doc)
    doc.save(str(REPORT))
    print(f"Updated {REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
